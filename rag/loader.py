"""Document loading utilities for RAG pipeline.

This module provides helpers to load documents from common file types
(PDF, TXT, DOCX) using LangChain loaders and to split them into overlapping
chunks using RecursiveCharacterTextSplitter.

Chunking rationale (why we split documents):
- Large documents are split into smaller chunks so that embedding models and
  vector search operate on meaningful, fixed-size contexts. This improves
  retrieval relevance and keeps token usage predictable.
- We use overlapping chunks to preserve context that crosses chunk boundaries
  (the overlap helps QA models that depend on context surrounding a token).

The default chunk size and overlap are chosen to balance context length and
efficient retrieval: chunk_size=800, chunk_overlap=100.
"""

from pathlib import Path
from typing import List, Union, Optional
import logging

# LangChain document types and splitters (modular package layout)
from langchain_core.documents import Document
from langchain_community.document_loaders import (
	PyPDFLoader,
	TextLoader,
	Docx2txtLoader,
	UnstructuredWordDocumentLoader,
	UnstructuredPDFLoader,
	UnstructuredPowerPointLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
# Additional optional dependencies
from bs4 import BeautifulSoup
import json
import pandas as pd
from lxml import etree

# Optional parsing libraries with safe fallbacks
try:
	import docx as python_docx  # python-docx
except Exception:  # pragma: no cover - optional dependency
	python_docx = None

try:
	import openpyxl  # xlsx support for pandas  # type: ignore[import]
except Exception:  # pragma: no cover - optional dependency
	openpyxl = None

logger = logging.getLogger(__name__)
logger.addHandler(logging.NullHandler())

# Provide safe local fallbacks when langchain splitters or Document are missing.
# These fallbacks ensure the loader works even when langchain is not installed
# (useful for local, offline environments). They implement a minimal subset
# of the RecursiveCharacterTextSplitter and Document API used by the rest
# of the codebase.
if Document is None:
	from dataclasses import dataclass

	@dataclass
	class Document:
		page_content: str
		metadata: dict = None


class _FallbackTextSplitter:
	"""A minimal text splitter compatible with langchain's API.

	This splitter chops text into character chunks with an overlap. It prefers
	splitting at paragraph boundaries when possible but falls back to fixed
	character sizes. It implements split_text(text) and split_documents(docs).
	"""
	def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100) -> None:
		self.chunk_size = int(chunk_size)
		self.chunk_overlap = int(chunk_overlap)

	def split_text(self, text: str) -> List[str]:
		if text is None:
			return []
		text = text.strip()
		if not text:
			return []

		chunks: List[str] = []
		start = 0
		L = len(text)
		while start < L:
			end = min(start + self.chunk_size, L)
			# prefer to cut at last double newline within range to preserve paragraphs
			sub = text[start:end]
			sep = sub.rfind("\n\n")
			if sep != -1 and sep > int(self.chunk_size * 0.33):
				end = start + sep
			chunk = text[start:end].strip()
			if chunk:
				chunks.append(chunk)
			if end >= L:
				break
			start = max(0, end - self.chunk_overlap)
		return chunks

	def split_documents(self, docs: List[Document]) -> List[Document]:
		out: List[Document] = []
		for d in docs:
			text = getattr(d, "page_content", None) or getattr(d, "text", None) or str(d)
			md = dict(getattr(d, "metadata", {}) or {})
			for c in self.split_text(text):
				out.append(Document(page_content=c, metadata=md))
		return out


def get_text_splitter(chunk_size: int = 800, chunk_overlap: int = 100):
	"""Return an object with the same interface as RecursiveCharacterTextSplitter.

	If LangChain's splitter is available, return an instance of it. Otherwise
	return an instance of the internal fallback. This keeps downstream code
	agnostic about the installed environment.
	"""
	if RecursiveCharacterTextSplitter is not None:
		return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
	return _FallbackTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)



def _choose_pdf_loader():
	"""Return an available PDF loader class from LangChain.

	We try a couple of common loaders to maximize compatibility across
	LangChain versions and optional dependencies.
	"""
	if PyPDFLoader is not None:
		return PyPDFLoader
	if UnstructuredPDFLoader is not None:
		return UnstructuredPDFLoader
	raise RuntimeError("No PDF loader is available. Install langchain and its PDF extras.")


def _choose_pptx_loader():
	"""Return an available PPTX loader class from LangChain if present."""
	if 'UnstructuredPowerPointLoader' in globals() and UnstructuredPowerPointLoader is not None:
		return UnstructuredPowerPointLoader
	# If not available, we explicitly state the dependency is missing.
	raise RuntimeError("No PPTX loader available. Install langchain with unstructured or the pptx extras.")


def _choose_docx_loader():
	"""Return an available DOCX loader class from LangChain.

	Prefers lightweight 'docx2txt' loader if present, falls back to
	unstructured loader if available.
	"""
	if Docx2txtLoader is not None:
		return Docx2txtLoader
	if UnstructuredWordDocumentLoader is not None:
		return UnstructuredWordDocumentLoader
	raise RuntimeError("No DOCX loader is available. Install langchain and its word/document extras.")


def _get_loader_for_path(path: Union[str, Path]):
	"""Return a loader instance or a text-extracting callable for the given file path.

	Supported extensions: .pdf, .txt, .docx, .pptx, .xml, .json, .html, .csv
	"""
	p = Path(path)
	if not p.exists():
		raise FileNotFoundError(f"File not found: {p}")

	suffix = p.suffix.lower()
	if suffix == ".pdf":
		LoaderCls = _choose_pdf_loader()
		def _pdf_loader():
			loader = LoaderCls(str(p))
			# Most LangChain PDF loaders implement `load()` -> List[Document]
			if hasattr(loader, "load"):
				return loader.load()
			# Fallback: some loaders may expose `load_and_split` or similar
			if hasattr(loader, "load_and_split"):
				return loader.load_and_split()
			# No supported loading method found on this loader class
			raise RuntimeError("PDF loader class does not implement 'load' or 'load_and_split'")
		return _pdf_loader
	if suffix == ".txt":
		if TextLoader is None:
			raise RuntimeError("TextLoader is not available. Ensure langchain is installed.")
		def _txt_loader():
			loader = TextLoader(str(p), encoding="utf-8")
			if hasattr(loader, "load"):
				return loader.load()
			# No supported loading method found on this loader class
			raise RuntimeError("Text loader class does not implement 'load'")
		return _txt_loader
	if suffix == ".docx":
		# Prefer LangChain loaders when available; otherwise fall back to python-docx
		try:
			LoaderCls = _choose_docx_loader()
			def _docx_loader():
				loader = LoaderCls(str(p))
				if hasattr(loader, "load"):
					return loader.load()
				if hasattr(loader, "load_and_split"):
					return loader.load_and_split()
				# No supported loading method found on this loader class
				raise RuntimeError("DOCX loader class does not implement 'load' or 'load_and_split'")
			return _docx_loader
		except Exception:
			# Fallback: use python-docx to extract paragraphs and tables
			if python_docx is None:
				raise RuntimeError("No DOCX loader is available. Install python-docx or langchain extras for DOCX support.")

			def _docx_fallback():
				doc = python_docx.Document(str(p))
				parts = []
				# paragraphs
				for para in doc.paragraphs:
					text = para.text.strip()
					if text:
						parts.append(text)
				# tables: convert rows to readable sentences preserving headers
				for table in doc.tables:
					# attempt to detect header row
					head = []
					if table.rows:
						for c in table.rows[0].cells:
							head.append(c.text.strip())
					for r in table.rows[1:]:
						cells = [c.text.strip() for c in r.cells]
						if any(cells):
							if any(head for head in head):
								pairs = [f"{h}: {v}" for h, v in zip(head, cells)]
								parts.append("; ".join(pairs))
				return "\n\n".join(parts)

			return _docx_fallback
	if suffix == ".pptx":
		LoaderCls = _choose_pptx_loader()
		def _pptx_loader():
			loader = LoaderCls(str(p))
			if hasattr(loader, "load"):
				return loader.load()
			if hasattr(loader, "load_and_split"):
				return loader.load_and_split()
			# No supported loading method found on this loader class
			raise RuntimeError("PPTX loader class does not implement 'load' or 'load_and_split'")
		return _pptx_loader
	if suffix == ".xml":
		# Use lxml to parse and extract text nodes
		def _xml_loader():
			parser = etree.parse(str(p))
			root = parser.getroot()
			texts = []
			for elem in root.iter():
				if elem.text and elem.text.strip():
					texts.append(elem.text.strip())
			return "\n\n".join(texts)
		return _xml_loader
	if suffix == ".json":
		def _json_loader():
			with open(p, "r", encoding="utf-8") as fh:
				obj = json.load(fh)
			return json.dumps(obj, indent=2, ensure_ascii=False)
		return _json_loader
	if suffix == ".html" or suffix == ".htm":
		def _html_loader():
			with open(p, "r", encoding="utf-8") as fh:
				soup = BeautifulSoup(fh, "html.parser")
				for s in soup(["script", "style"]):
					s.delete()
				return soup.get_text(separator="\n")
		return _html_loader
	if suffix == ".csv":
		def _csv_loader():
			# Read as CSV and convert rows to short sentences
			df = pd.read_csv(p)
			rows = []
			for _, r in df.iterrows():
				parts = [f"{c}: {r[c]}" for c in df.columns]
				rows.append("; ".join(parts))
			return "\n\n".join(rows)
		return _csv_loader

	if suffix in (".xls", ".xlsx"):
		def _excel_loader():
			# Read all sheets and convert into readable narratives preserving headers
			df_dict = pd.read_excel(p, sheet_name=None)
			parts = []
			for sheet_name, df in df_dict.items():
				parts.append(f"SHEET: {sheet_name}")
				if df.empty:
					parts.append("(empty sheet)")
				else:
					for _, r in df.iterrows():
						cells = [f"{c}: {r[c]}" for c in df.columns]
						parts.append("; ".join(cells))
			return "\n\n".join(parts)
		return _excel_loader
def load_and_split_documents(
	paths: Union[str, Path, List[Union[str, Path]]],
	chunk_size: int = 800,
	chunk_overlap: int = 100,
) -> List[Document]:
	"""Load documents (paths or raw text) and split them into chunks.

	This function normalizes input so agents can operate on text-only
	Document objects regardless of source format. If ``paths`` are file
	paths, each supported file type is converted to text using safe,
	format-specific extractors. If a string is provided that is not a
	filesystem path, it is treated as raw document text.
	"""
	if RecursiveCharacterTextSplitter is None and _FallbackTextSplitter is None:
		raise RuntimeError(
			"RecursiveCharacterTextSplitter is not available. Ensure langchain is installed or use the fallback splitter."
		)

	# Normalize to list of sources
	if isinstance(paths, (str, Path)):
		sources = [paths]
	else:
		sources = list(paths)

	all_docs: List[Document] = []
	splitter = get_text_splitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

	for src in sources:
		p = Path(src)
		raw_texts: List[str] = []
		if isinstance(src, (str, Path)) and p.exists():
			loader = _get_loader_for_path(p)
			result = loader()
			# If loader returned a list of Document objects, split them directly
			if isinstance(result, list) and result and isinstance(result[0], Document):
				chunks = splitter.split_documents(result)
				all_docs.extend(chunks)
			else:
				# Treat result as plain text (string)
				text = str(result)
				if text and text.strip():
					doc = Document(page_content=text, metadata={"source": str(p)})
					chunks = splitter.split_documents([doc])
					all_docs.extend(chunks)
		else:
			# treat src as raw text
			raw_texts.append(str(src))

		for t in raw_texts:
			if not t or not t.strip():
				continue
			doc = Document(page_content=t, metadata={"source": None})
			chunks = splitter.split_documents([doc])
			all_docs.extend(chunks)

	logger.info("Loaded and split %d chunks from %d sources", len(all_docs), len(sources))
	return all_docs


# Backwards compatibility alias
load_and_split = load_and_split_documents


if __name__ == "__main__":
	# Quick CLI for manual testing. Not intended for production use, but
	# convenient for developers.
	import argparse

	parser = argparse.ArgumentParser(description="Load and chunk documents for RAG.")
	parser.add_argument("paths", nargs="+", help="Files to load (pdf, txt, docx)")
	parser.add_argument("--chunk-size", type=int, default=800)
	parser.add_argument("--overlap", type=int, default=100)
	args = parser.parse_args()

	logging.basicConfig(level=logging.INFO)
	docs = load_and_split(args.paths, chunk_size=args.chunk_size, chunk_overlap=args.overlap)
	print(f"Produced {len(docs)} document chunks")

